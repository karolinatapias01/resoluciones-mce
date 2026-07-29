# -*- coding: utf-8 -*-
"""
GENERADOR DE RESOLUCIONES MCE - Fondo Rotatorio de Estupefacientes (Santander)
================================================================================
Lee un archivo JSON (exportado del formulario) y genera un DOCX por cada
registro, usando las plantillas que conservan membrete, encabezado y pie.

USO:
    python generar_resoluciones.py datos.json

Genera los archivos en la carpeta ./resoluciones_generadas/
"""
import sys, os, json, re
from copy import deepcopy
from docx import Document

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
def seleccionar_plantilla(tramite, est):
    if tramite == 'AMPL':
        if est == 'MAY':
            return 'plantillas/TPL_AMPL_MAYORISTA.docx'
        elif est in ('IPS-P', 'IPS-G'):
            return 'plantillas/TPL_AMPL_IPS.docx'
        return 'plantillas/TPL_AMPL_DROGUERIA_MINORISTA_IPS.docx'
    elif tramite == 'MOD-DT':
        if est in ('IPS-P', 'IPS-G'):
            return 'plantillas/TPL_MOD_DT_IPS.docx'
        return 'plantillas/TPL_MOD_DT.docx'
    elif tramite == 'MOD-VAR':
        if est in ('IPS-P', 'IPS-G'):
            return 'plantillas/TPL_MOD_VAR_IPS.docx'
        return 'plantillas/TPL_MOD_VAR.docx'
    elif tramite == 'INSC':
        if est in ('IPS-P', 'IPS-G'):
            return 'plantillas/TPL_INSC_IPS.docx'
        return 'plantillas/TPL_INSC.docx'
    elif tramite == 'RENOV':
        if est in ('IPS-P', 'IPS-G'):
            return 'plantillas/TPL_RENOV_IPS.docx'
        elif est == 'MAY':
            return 'plantillas/TPL_RENOV_MAY.docx'
        return 'plantillas/TPL_RENOV_DROG.docx'
    return {
        'ACLARA':  'plantillas/TPL_ACLARA.docx',
        'CANCEL':  'plantillas/TPL_CANCEL.docx',
    }.get(tramite)
TIPO_EST_DESC = {
    'DROG':'', 'IPS-P':'', 'IPS-G':'', 'MAY':'mayorista', 'VET':'veterinario'
}

def repl_text(text, ctx):
    for k, v in ctx.items():
        text = text.replace('{{'+k+'}}', str(v)).replace('{{ '+k+' }}', str(v))
    return text

def process_paragraph(p, ctx):
    # Reemplazo run-por-run: conserva la negrilla/formato de cada run.
    # Como cada marcador {{x}} ya vive completo dentro de un solo run
    # (así se construyeron las plantillas), basta reemplazar en cada run.
    for r in p.runs:
        if '{{' in r.text:
            r.text = repl_text(r.text, ctx)
    # Por si algún marcador quedó partido entre runs, pasada de seguridad:
    full = ''.join(r.text for r in p.runs)
    if '{{' in full and '}}' in full:
        new = repl_text(full, ctx)
        if new != full and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ''

def fill_row(tr, med):
    cells = tr.findall(f'.//{W}tc')
    n = len(cells)
    if n >= 5:
        # Encabezado con columnas combinadas: NOMBRE | CONC | CONC | FORMA | FORMA
        valores = [med['nombre'], med['concentracion'], med['concentracion'], med['forma'], med['forma']]
    elif n == 4:
        valores = [med['nombre'], med['concentracion'], med['forma'], med['forma']]
    else:
        valores = [med['nombre'], med['concentracion'], med['forma']]
    for tc, val in zip(cells, valores):
        ts = tc.findall(f'.//{W}t')
        if ts:
            ts[0].text = val
            for extra in ts[1:]:
                extra.text = ''

def llenar_tabla(tbl, mce):
    # Guardar una fila plantilla (la primera de datos) y borrar TODAS las de datos viejas
    tmpl = deepcopy(tbl.rows[1]._tr)
    filas_datos = tbl.rows[1:]  # todas menos el encabezado
    for row in filas_datos:
        row._tr.getparent().remove(row._tr)
    # Insertar filas nuevas basadas en la plantilla
    header_tr = tbl.rows[0]._tr
    prev = header_tr
    for med in mce:
        new_tr = deepcopy(tmpl)
        fill_row(new_tr, med)
        prev.addnext(new_tr)
        prev = new_tr

def es_tabla_mce(tbl):
    header = [c.text.strip().upper().replace('É', 'E') for c in tbl.rows[0].cells]
    return 'NOMBRE GENERICO' in header

def generar(registro, salida):
    tramite = registro['tramite']
    plantilla = seleccionar_plantilla(tramite, registro.get('establecimiento',''))
    if not plantilla or not os.path.exists(plantilla):
        print(f"  ⚠ Sin plantilla para {tramite}, se omite.")
        return False

    # Construir contexto de texto
    ctx = dict(registro)
    # Respetar lo que envia el formulario; solo completar si viene vacio
    # Articulo el/la derivado del tratamiento
    if not ctx.get('articulo'):
        ctx['articulo'] = 'la' if str(registro.get('tratamiento','')).strip().startswith('la ') else 'el'
    if not ctx.get('dt_articulo'):
        ctx['dt_articulo'] = 'la' if str(registro.get('dt_tratamiento','')).strip().startswith('la ') else 'el'
    if not ctx.get('tipo_est_desc'):
        ctx['tipo_est_desc'] = TIPO_EST_DESC.get(registro.get('establecimiento',''), '')
    # Canal: solo armarlo si el formulario no lo mando ya
    if not ctx.get('canal_texto'):
        canal = registro.get('canal', '')
        if canal == 'Plataforma FOREST':
            ctx['canal_texto'] = f"vía plataforma FOREST con N° de proceso {registro.get('forest_proceso','')}"
        elif canal:
            ctx['canal_texto'] = f"vía {canal.lower()}"
        else:
            ctx['canal_texto'] = ''
    # Tratamiento por calidad (heurística simple)
    ctx.setdefault('tratamiento', 'el/la señor(a)')

    # Parsear MCE
    mce = []
    for linea in registro.get('mce', '').strip().split('\n'):
        if '|' in linea:
            partes = [x.strip() for x in linea.split('|')]
            if len(partes) >= 3:
                mce.append({'nombre': partes[0], 'concentracion': partes[1], 'forma': partes[2]})

    mono = []
    for linea in str(registro.get('mce_monopolio', '')).strip().split('\n'):
        if '|' in linea:
            partes = [x.strip() for x in linea.split('|')]
            if len(partes) >= 3:
                mono.append({'nombre': partes[0], 'concentracion': partes[1], 'forma': partes[2]})

    d = Document(plantilla)
    for p in d.paragraphs:
        process_paragraph(p, ctx)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p, ctx)
    # Clasificar cada tabla MCE como monopolio o control especial segun el parrafo que la precede
    from docx.oxml.ns import qn
    def texto_antes(tbl):
        el = tbl._tbl.getprevious()
        encontrados = []
        saltos = 0
        while el is not None and saltos < 8 and len(encontrados) < 3:
            if el.tag == qn('w:p'):
                t = ''.join(n.text or '' for n in el.iter(qn('w:t'))).strip()
                if t:
                    encontrados.append(t)
            el = el.getprevious(); saltos += 1
        return ' '.join(encontrados).upper()

    for tbl in d.tables:
        if not es_tabla_mce(tbl):
            continue
        antes = texto_antes(tbl)
        es_mono = ('MONOPOLIO' in antes) or ('PROHIBIDA' in antes)
        datos = mono if es_mono else mce
        if datos:
            llenar_tabla(tbl, datos)
        else:
            # Sin datos: dejar solo el encabezado
            for row in list(tbl.rows[1:]):
                row._tr.getparent().remove(row._tr)

    # Lista de documentos numerada — SOLO en el considerando SEGUNDO (allegó documentos)
    docs_raw = registro.get('documentos', '')
    docs = [l.strip() for l in str(docs_raw).split('\n') if l.strip()]
    if docs:
        # Localizar el parrafo del considerando SEGUNDO donde se allegan documentos
        ini_zona = None
        fin_zona = None
        for i, p in enumerate(d.paragraphs):
            t = p.text.strip()
            tu = t.upper()
            if ini_zona is None and tu.startswith('SEGUNDO') and ('ALLEG' in tu or 'DOCUMENTO' in tu):
                ini_zona = i
                continue
            if ini_zona is not None:
                # la zona termina en el siguiente considerando (TERCERO/CUARTO...) o ARTICULO
                if tu.startswith(('TERCERO', 'CUARTO', 'QUINTO', 'SEXTO', 'ARTÍCULO', 'ARTICULO', 'RESUELVE')):
                    fin_zona = i
                    break
        # Recolectar los indices de items numerados SOLO dentro de esa zona
        idxs = []
        if ini_zona is not None:
            tope = fin_zona if fin_zona is not None else len(d.paragraphs)
            for i in range(ini_zona + 1, tope):
                t = d.paragraphs[i].text.strip()
                if len(t) > 2 and t[0].isdigit() and t[1] in '.)':
                    idxs.append(i)
        if idxs:
            base = d.paragraphs[idxs[0]]
            # Primero rellenar los que se conservan
            for j, txt in enumerate(docs, 1):
                if j - 1 < len(idxs):
                    par = d.paragraphs[idxs[j-1]]
                    if par.runs:
                        par.runs[0].text = f"{j}. {txt}"
                        for r in par.runs[1:]:
                            r.text = ''
                else:
                    from copy import deepcopy as _dc
                    nt = _dc(base._p)
                    base._p.getparent().insert(list(base._p.getparent()).index(base._p) + j, nt)
            # Borrar sobrantes de ATRAS hacia ADELANTE para no desincronizar
            if len(idxs) > len(docs):
                for k in range(len(idxs) - 1, len(docs) - 1, -1):
                    pr = d.paragraphs[idxs[k]]
                    pr._p.getparent().remove(pr._p)
    d.save(salida)
    return True

def main():
    if len(sys.argv) < 2:
        print("USO: python generar_resoluciones.py datos.json")
        sys.exit(1)
    with open(sys.argv[1], encoding='utf-8') as f:
        registros = json.load(f)
    if isinstance(registros, dict):
        registros = [registros]
    os.makedirs('resoluciones_generadas', exist_ok=True)
    ok = 0
    for i, reg in enumerate(registros, 1):
        nombre = re.sub(r'[^\w\-]', '_', reg.get('nombre_est', f'registro_{i}'))
        salida = f'resoluciones_generadas/{i:02d}_{reg["tramite"]}_{nombre}.docx'
        if generar(reg, salida):
            print(f"  ✓ {salida}")
            ok += 1
    print(f"\n{ok} de {len(registros)} resoluciones generadas en ./resoluciones_generadas/")

if __name__ == '__main__':
    main()
